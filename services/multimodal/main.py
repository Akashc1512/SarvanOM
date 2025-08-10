"""
Multimodal File Processing Service

Handles upload and processing of various file types:
- Images (PNG, JPG, JPEG, GIF, WebP)
- Videos (MP4, AVI, MOV, WebM)
- Documents (PDF, DOC, DOCX, TXT, MD)
- Audio (MP3, WAV, OGG)
- Archives (ZIP, RAR, 7Z)
"""

import os
import time
import uuid
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

from fastapi import FastAPI, File, UploadFile, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
import aiofiles
from PIL import Image
import pymupdf  # PyMuPDF for PDF processing
import docx
import moviepy.editor as mp
import whisper
import zipfile
import rarfile

from shared.core.config import get_central_config
from shared.core.logging import get_logger
from shared.embeddings.local_embedder import embed_texts
from shared.vectorstores.vector_store_service import VectorDocument
from shared.contracts.query import RetrievalIndexRequest

logger = get_logger(__name__)

# Load configuration
config = get_central_config()

app = FastAPI(
    title="SarvanOM Multimodal Processing Service",
    version="1.0.0",
    description="Handle multimodal file uploads and content extraction",
)

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=config.cors_origins,
    allow_credentials=bool(config.cors_credentials),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Trusted hosts
try:
    from shared.core.api.config import get_settings
    _settings = get_settings()
    _allowed_hosts = _settings.trusted_hosts or ["*"]
except Exception:
    _allowed_hosts = ["*"]

app.add_middleware(TrustedHostMiddleware, allowed_hosts=_allowed_hosts)

# File type configurations
SUPPORTED_IMAGE_TYPES = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff"}
SUPPORTED_VIDEO_TYPES = {".mp4", ".avi", ".mov", ".webm", ".mkv", ".flv"}
SUPPORTED_DOCUMENT_TYPES = {".pdf", ".doc", ".docx", ".txt", ".md", ".rtf"}
SUPPORTED_AUDIO_TYPES = {".mp3", ".wav", ".ogg", ".m4a", ".flac"}
SUPPORTED_ARCHIVE_TYPES = {".zip", ".rar", ".7z", ".tar", ".gz"}

MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
UPLOAD_DIR = Path("data/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Initialize Whisper model for audio transcription
whisper_model = None

async def get_whisper_model():
    global whisper_model
    if whisper_model is None:
        whisper_model = whisper.load_model("base")
    return whisper_model


@app.get("/health")
async def health():
    return {
        "service": "multimodal",
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "supported_types": {
            "images": list(SUPPORTED_IMAGE_TYPES),
            "videos": list(SUPPORTED_VIDEO_TYPES),
            "documents": list(SUPPORTED_DOCUMENT_TYPES),
            "audio": list(SUPPORTED_AUDIO_TYPES),
            "archives": list(SUPPORTED_ARCHIVE_TYPES),
        }
    }


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """Upload and process a multimodal file."""
    try:
        # Validate file size
        if file.size and file.size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
            )

        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_extension = Path(file.filename or "").suffix.lower()
        temp_filename = f"{file_id}{file_extension}"
        temp_path = UPLOAD_DIR / temp_filename

        # Save uploaded file
        async with aiofiles.open(temp_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        logger.info(f"File uploaded: {file.filename} -> {temp_path}")

        # Process file based on type
        extracted_content = await process_file(temp_path, file_extension, file.filename)

        # Index content in vector database if text was extracted
        if extracted_content.get("text"):
            await index_content(file_id, extracted_content, file.filename)

        # Clean up temporary file (optional - you might want to keep for caching)
        # os.unlink(temp_path)

        return {
            "file_id": file_id,
            "filename": file.filename,
            "file_type": get_file_type(file_extension),
            "processed": True,
            "extracted_content": extracted_content,
            "indexed": bool(extracted_content.get("text")),
            "timestamp": datetime.now().isoformat(),
        }

    except Exception as e:
        logger.error(f"File upload failed: {e}")
        # Clean up on error
        if 'temp_path' in locals() and temp_path.exists():
            os.unlink(temp_path)
        raise HTTPException(status_code=500, detail=f"File processing failed: {str(e)}")


async def process_file(file_path: Path, file_extension: str, original_filename: str) -> Dict[str, Any]:
    """Process file based on its type and extract content."""
    
    file_type = get_file_type(file_extension)
    extracted_content = {
        "file_type": file_type,
        "file_extension": file_extension,
        "original_filename": original_filename,
        "text": "",
        "metadata": {},
        "processed_at": datetime.now().isoformat(),
    }

    try:
        if file_type == "image":
            extracted_content.update(await process_image(file_path))
        elif file_type == "video":
            extracted_content.update(await process_video(file_path))
        elif file_type == "document":
            extracted_content.update(await process_document(file_path, file_extension))
        elif file_type == "audio":
            extracted_content.update(await process_audio(file_path))
        elif file_type == "archive":
            extracted_content.update(await process_archive(file_path))
        else:
            extracted_content["text"] = f"Unsupported file type: {file_extension}"

    except Exception as e:
        logger.error(f"Error processing {file_type} file {file_path}: {e}")
        extracted_content["error"] = str(e)
        extracted_content["text"] = f"Error processing {file_type}: {str(e)}"

    return extracted_content


def get_file_type(file_extension: str) -> str:
    """Determine file type category from extension."""
    if file_extension in SUPPORTED_IMAGE_TYPES:
        return "image"
    elif file_extension in SUPPORTED_VIDEO_TYPES:
        return "video"
    elif file_extension in SUPPORTED_DOCUMENT_TYPES:
        return "document"
    elif file_extension in SUPPORTED_AUDIO_TYPES:
        return "audio"
    elif file_extension in SUPPORTED_ARCHIVE_TYPES:
        return "archive"
    else:
        return "unknown"


async def process_image(file_path: Path) -> Dict[str, Any]:
    """Process image files - extract metadata and generate description."""
    try:
        with Image.open(file_path) as img:
            width, height = img.size
            format_name = img.format
            mode = img.mode
            
            # Basic image analysis
            text_content = f"Image file: {format_name} format, {width}x{height} pixels, {mode} color mode."
            
            # TODO: Add image captioning with CLIP or similar model
            # For now, provide basic metadata as text
            
            return {
                "text": text_content,
                "metadata": {
                    "width": width,
                    "height": height,
                    "format": format_name,
                    "mode": mode,
                    "has_transparency": mode in ("RGBA", "LA") or "transparency" in img.info,
                }
            }
    except Exception as e:
        return {"text": f"Error processing image: {str(e)}", "metadata": {}}


async def process_video(file_path: Path) -> Dict[str, Any]:
    """Process video files - extract metadata and transcribe audio."""
    try:
        # Extract video metadata
        clip = mp.VideoFileClip(str(file_path))
        duration = clip.duration
        fps = clip.fps
        size = clip.size
        
        text_content = f"Video file: {duration:.2f} seconds duration, {fps} FPS, {size[0]}x{size[1]} resolution."
        
        # Extract audio for transcription
        if clip.audio is not None:
            # Extract audio to temporary file
            audio_path = file_path.with_suffix('.wav')
            clip.audio.write_audiofile(str(audio_path), verbose=False, logger=None)
            
            # Transcribe audio
            model = await get_whisper_model()
            result = model.transcribe(str(audio_path))
            transcription = result["text"]
            
            text_content += f"\n\nVideo transcription: {transcription}"
            
            # Clean up audio file
            os.unlink(audio_path)
        
        clip.close()
        
        return {
            "text": text_content,
            "metadata": {
                "duration": duration,
                "fps": fps,
                "width": size[0],
                "height": size[1],
                "has_audio": clip.audio is not None,
            }
        }
    except Exception as e:
        return {"text": f"Error processing video: {str(e)}", "metadata": {}}


async def process_document(file_path: Path, file_extension: str) -> Dict[str, Any]:
    """Process document files - extract text content."""
    try:
        text_content = ""
        metadata = {}
        
        if file_extension == ".pdf":
            # Use PyMuPDF for PDF processing
            doc = pymupdf.open(str(file_path))
            text_parts = []
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_parts.append(page.get_text())
            text_content = "\n".join(text_parts)
            metadata = {
                "pages": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
            }
            doc.close()
            
        elif file_extension in [".doc", ".docx"]:
            # Use python-docx for Word documents
            doc = docx.Document(str(file_path))
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            text_content = "\n".join(text_parts)
            metadata = {"paragraphs": len(doc.paragraphs)}
            
        elif file_extension in [".txt", ".md", ".rtf"]:
            # Plain text files
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = await f.read()
            metadata = {"lines": len(text_content.splitlines())}
        
        return {
            "text": text_content,
            "metadata": metadata
        }
    except Exception as e:
        return {"text": f"Error processing document: {str(e)}", "metadata": {}}


async def process_audio(file_path: Path) -> Dict[str, Any]:
    """Process audio files - transcribe to text."""
    try:
        # Use Whisper for transcription
        model = await get_whisper_model()
        result = model.transcribe(str(file_path))
        
        text_content = f"Audio transcription: {result['text']}"
        
        return {
            "text": text_content,
            "metadata": {
                "language": result.get("language", "unknown"),
                "duration": result.get("duration", 0),
            }
        }
    except Exception as e:
        return {"text": f"Error processing audio: {str(e)}", "metadata": {}}


async def process_archive(file_path: Path) -> Dict[str, Any]:
    """Process archive files - list contents and extract text from supported files."""
    try:
        text_parts = []
        file_list = []
        
        if file_path.suffix.lower() == ".zip":
            with zipfile.ZipFile(file_path, 'r') as archive:
                file_list = archive.namelist()
                # Extract and process text files from archive
                for file_name in file_list[:10]:  # Limit to first 10 files
                    if any(file_name.lower().endswith(ext) for ext in ['.txt', '.md', '.py', '.js', '.json']):
                        try:
                            content = archive.read(file_name).decode('utf-8', errors='ignore')
                            text_parts.append(f"File: {file_name}\n{content[:1000]}...")  # First 1000 chars
                        except:
                            continue
                            
        elif file_path.suffix.lower() == ".rar":
            with rarfile.RarFile(file_path, 'r') as archive:
                file_list = archive.namelist()
                # RAR extraction is more complex, just list files for now
        
        text_content = f"Archive contains {len(file_list)} files.\n"
        if text_parts:
            text_content += "\n".join(text_parts)
        
        return {
            "text": text_content,
            "metadata": {
                "file_count": len(file_list),
                "files": file_list[:20],  # First 20 filenames
            }
        }
    except Exception as e:
        return {"text": f"Error processing archive: {str(e)}", "metadata": {}}


async def index_content(file_id: str, content: Dict[str, Any], filename: str):
    """Index extracted content in vector database."""
    try:
        text = content.get("text", "")
        if not text or len(text.strip()) < 10:
            return
        
        # Create embeddings
        embeddings = embed_texts([text])
        
        # Create vector document
        document = VectorDocument(
            id=file_id,
            text=text,
            embedding=embeddings[0],
            metadata={
                "filename": filename,
                "file_type": content.get("file_type"),
                "processed_at": content.get("processed_at"),
                "source": "user_upload",
                **content.get("metadata", {})
            }
        )
        
        # Index in vector store (send to retrieval service)
        import httpx
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                indexing_payload = RetrievalIndexRequest(
                    ids=[file_id],
                    texts=[text],
                    metadatas=[document.metadata]
                )
                response = await client.post(
                    f"{config.search_service_url}/index",
                    json=indexing_payload.dict()
                )
                response.raise_for_status()
                logger.info(f"Successfully indexed content for file {file_id}")
        except Exception as e:
            logger.error(f"Failed to index content: {e}")
            
    except Exception as e:
        logger.error(f"Error indexing content: {e}")


@app.get("/files/{file_id}")
async def get_file_info(file_id: str):
    """Get information about a processed file."""
    # This would typically query a database for file metadata
    # For now, return a placeholder
    return {
        "file_id": file_id,
        "status": "processed",
        "message": "File information retrieval not implemented yet"
    }


@app.delete("/files/{file_id}")
async def delete_file(file_id: str):
    """Delete a processed file and its content from vector database."""
    try:
        # Remove file from disk
        file_pattern = UPLOAD_DIR / f"{file_id}.*"
        for file_path in UPLOAD_DIR.glob(f"{file_id}.*"):
            os.unlink(file_path)
        
        # TODO: Remove from vector database
        # This would require implementing a delete endpoint in the retrieval service
        
        return {
            "file_id": file_id,
            "deleted": True,
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete file: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("services.multimodal.main:app", host="0.0.0.0", port=8006, reload=True)
